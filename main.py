import customtkinter as ctk
from tkinter import filedialog, messagebox
from PIL import Image, ImageTk
import numpy as np
import PyPDF2
import base64
from docx import Document

# ---------- IMAGE STEGANOGRAPHY ----------
def hide_image(cover_image_path, secret_image_path, output_image_path):
    try:
        cover_image = Image.open(cover_image_path).convert("RGB")
        secret_image = Image.open(secret_image_path).convert("L")
        secret_image = secret_image.resize(cover_image.size)
        cover_array = np.array(cover_image)
        secret_array = np.array(secret_image)
        secret_array = np.stack((secret_array,) * 3, axis=-1)
        cover_array = (cover_array & 0xF0) | (secret_array >> 4)
        encoded_image = Image.fromarray(cover_array, "RGB")
        encoded_image.save(output_image_path)
        messagebox.showinfo("Success", f"Hidden image saved at {output_image_path}")
    except Exception as e:
        messagebox.showerror("Error", f"Failed: {e}")

def extract_image(encoded_image_path, output_secret_image_path):
    try:
        encoded_image = Image.open(encoded_image_path).convert("RGB")
        encoded_array = np.array(encoded_image)
        secret_array = (encoded_array & 0x0F) << 4
        secret_image = Image.fromarray(secret_array[:, :, 0], "L")
        secret_image.save(output_secret_image_path)
        messagebox.showinfo("Success", f"Extracted secret saved at {output_secret_image_path}")
    except Exception as e:
        messagebox.showerror("Error", f"Failed: {e}")

# ---------- PDF ----------
def embed_pdf_in_another(pdf_to_hide_path, base_pdf_path, output_pdf_path):
    try:
        binary_data = open(pdf_to_hide_path, 'rb').read()
        encoded_data = base64.b64encode(binary_data).decode('utf-8')
        reader = PyPDF2.PdfReader(base_pdf_path)
        writer = PyPDF2.PdfWriter()
        for page in reader.pages:
            writer.add_page(page)
        writer.add_metadata({'/HiddenData': encoded_data})
        with open(output_pdf_path, 'wb') as f:
            writer.write(f)
        messagebox.showinfo("Success", f"PDF embedded at {output_pdf_path}")
    except Exception as e:
        messagebox.showerror("Error", str(e))

def extract_pdf(pdf_path, output_pdf_path):
    try:
        reader = PyPDF2.PdfReader(pdf_path)
        meta = reader.metadata
        if '/HiddenData' in meta:
            data = base64.b64decode(meta['/HiddenData'])
            open(output_pdf_path, 'wb').write(data)
            messagebox.showinfo("Success", f"Extracted PDF at {output_pdf_path}")
        else:
            messagebox.showwarning("No Data", "No hidden data found.")
    except Exception as e:
        messagebox.showerror("Error", str(e))

# ---------- DOCX ----------
def embed_docx_in_another(docx_to_hide_path, base_docx_path, output_docx_path):
    try:
        data = base64.b64encode(open(docx_to_hide_path, 'rb').read()).decode('utf-8')
        base = Document(base_docx_path)
        para = base.add_paragraph()
        run = para.add_run(data)
        run.font.hidden = True
        base.save(output_docx_path)
        messagebox.showinfo("Success", f"Embedded DOCX saved at {output_docx_path}")
    except Exception as e:
        messagebox.showerror("Error", str(e))

def extract_docx(docx_path, output_path):
    try:
        doc = Document(docx_path)
        for p in doc.paragraphs:
            for r in p.runs:
                if r.font.hidden:
                    open(output_path, 'wb').write(base64.b64decode(r.text))
                    messagebox.showinfo("Success", f"Extracted DOCX at {output_path}")
                    return
        messagebox.showwarning("No Data", "No hidden DOCX data found.")
    except Exception as e:
        messagebox.showerror("Error", str(e))

# ---------- UI ----------
ctk.set_appearance_mode("dark")
ctk.set_default_color_theme("blue")

class StegApp(ctk.CTk):
    def __init__(self):
        super().__init__()
        self.title("Steganography Tool")
        self.geometry("950x600")
        self.resizable(False, False)

        # Sidebar
        self.sidebar = ctk.CTkFrame(self, width=200)
        self.sidebar.pack(side="left", fill="y", padx=10, pady=10)
        ctk.CTkLabel(self.sidebar, text="🔐 Steganography Tool", font=("Arial", 16, "bold")).pack(pady=15)
        ctk.CTkButton(self.sidebar, text="Image Steganography", command=self.show_image_frame).pack(pady=10)
        ctk.CTkButton(self.sidebar, text="PDF Steganography", command=self.show_pdf_frame).pack(pady=10)
        ctk.CTkButton(self.sidebar, text="DOCX Steganography", command=self.show_docx_frame).pack(pady=10)

        # Logo Display
        img = Image.open("logo.png")
        img = img.resize((160, 120))
        self.logo = ImageTk.PhotoImage(img)
        ctk.CTkLabel(self.sidebar, image=self.logo, text="").pack(side="bottom", pady=10)

        # Main Area
        self.main_frame = ctk.CTkFrame(self, corner_radius=15)
        self.main_frame.pack(side="right", fill="both", expand=True, padx=10, pady=10)
        self.default_screen()

    def clear_main(self):
        for widget in self.main_frame.winfo_children():
            widget.destroy()

    def default_screen(self):
        self.clear_main()
        ctk.CTkLabel(self.main_frame, text="Welcome to the Steganography Tool", font=("Arial", 22, "bold")).pack(pady=40)
        ctk.CTkLabel(self.main_frame, text="Hide or extract secret data securely from images, PDFs, and DOCX files.",
                     font=("Arial", 14)).pack(pady=10)

    def show_image_frame(self):
        self.clear_main()
        ctk.CTkLabel(self.main_frame, text="🖼️ Image Steganography", font=("Arial", 18, "bold")).pack(pady=10)
        ctk.CTkButton(self.main_frame, text="Hide Secret in Image", command=self.hide_img_ui).pack(pady=10)
        ctk.CTkButton(self.main_frame, text="Extract Secret Image", command=self.extract_img_ui).pack(pady=10)

    def show_pdf_frame(self):
        self.clear_main()
        ctk.CTkLabel(self.main_frame, text="📄 PDF Steganography", font=("Arial", 18, "bold")).pack(pady=10)
        ctk.CTkButton(self.main_frame, text="Hide PDF", command=self.hide_pdf_ui).pack(pady=10)
        ctk.CTkButton(self.main_frame, text="Extract PDF", command=self.extract_pdf_ui).pack(pady=10)

    def show_docx_frame(self):
        self.clear_main()
        ctk.CTkLabel(self.main_frame, text="📝 DOCX Steganography", font=("Arial", 18, "bold")).pack(pady=10)
        ctk.CTkButton(self.main_frame, text="Hide DOCX", command=self.hide_docx_ui).pack(pady=10)
        ctk.CTkButton(self.main_frame, text="Extract DOCX", command=self.extract_docx_ui).pack(pady=10)

    # ---------- Actions ----------
    def hide_img_ui(self):
        cover = filedialog.askopenfilename(title="Select Cover Image")
        secret = filedialog.askopenfilename(title="Select Secret Image")
        out = filedialog.asksaveasfilename(defaultextension=".png")
        if cover and secret and out:
            hide_image(cover, secret, out)

    def extract_img_ui(self):
        enc = filedialog.askopenfilename(title="Select Encoded Image")
        out = filedialog.asksaveasfilename(defaultextension=".png")
        if enc and out:
            extract_image(enc, out)

    def hide_pdf_ui(self):
        pdf1 = filedialog.askopenfilename(title="Select PDF to Hide")
        pdf2 = filedialog.askopenfilename(title="Select Base PDF")
        out = filedialog.asksaveasfilename(defaultextension=".pdf")
        if pdf1 and pdf2 and out:
            embed_pdf_in_another(pdf1, pdf2, out)

    def extract_pdf_ui(self):
        enc = filedialog.askopenfilename(title="Select Encoded PDF")
        out = filedialog.asksaveasfilename(defaultextension=".pdf")
        if enc and out:
            extract_pdf(enc, out)

    def hide_docx_ui(self):
        d1 = filedialog.askopenfilename(title="Select DOCX to Hide")
        d2 = filedialog.askopenfilename(title="Select Base DOCX")
        out = filedialog.asksaveasfilename(defaultextension=".docx")
        if d1 and d2 and out:
            embed_docx_in_another(d1, d2, out)

    def extract_docx_ui(self):
        enc = filedialog.askopenfilename(title="Select Encoded DOCX")
        out = filedialog.asksaveasfilename(defaultextension=".docx")
        if enc and out:
            extract_docx(enc, out)


if __name__ == "__main__":
    app = StegApp()
    app.mainloop()
