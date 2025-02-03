import tkinter as tk
from tkinter import filedialog, messagebox, ttk
from PIL import Image
import numpy as np
import PyPDF2
import base64
import os
from docx import Document



def hide_image(cover_image_path, secret_image_path, output_image_path):
    try:
        cover_image = Image.open(cover_image_path).convert("RGB")
        secret_image = Image.open(secret_image_path).convert("L")
        secret_image = secret_image.resize(cover_image.size)

        cover_array = np.array(cover_image)
        secret_array = np.array(secret_image)
        secret_array = np.stack((secret_array,) * 3, axis=-1)

        cover_array = (cover_array & 0xF0) | (secret_array >> 4)
        encoded_image = Image.fromarray(cover_array, "RGB")
        encoded_image.save(output_image_path)
        messagebox.showinfo("Success", f"Image saved with hidden content at {output_image_path}")
    except Exception as e:
        messagebox.showerror("Error", f"Failed to hide image: {e}")

def extract_image(encoded_image_path, output_secret_image_path):
    try:
        encoded_image = Image.open(encoded_image_path).convert("RGB")
        encoded_array = np.array(encoded_image)
        secret_array = (encoded_array & 0x0F) << 4
        secret_image = Image.fromarray(secret_array[:, :, 0], "L")
        secret_image.save(output_secret_image_path)
        messagebox.showinfo("Success", f"Secret image extracted and saved at {output_secret_image_path}")
    except Exception as e:
        messagebox.showerror("Error", f"Failed to extract image: {e}")
# ------------------------ IMAGE PREDICTION LOGIC ------------------------

def predict_hidden_content(image_path):
    try:
        # Load the image and analyze pixel differences
        img = Image.open(image_path)
        img_array = np.array(img)

        # Simple prediction: Look for anomalies (this can be replaced with a trained model for real use cases)
        anomaly_score = np.sum(np.abs(img_array % 2))  # Example heuristic: check for low or high odd/even pixel values

        if anomaly_score > 100000:  # Threshold for detecting hidden content (just an example)
            messagebox.showinfo("Prediction", "Hidden content detected in the image!")
        else:
            messagebox.showinfo("Prediction", "No hidden content detected.")
    except Exception as e:
        messagebox.showerror("Error", f"Failed to predict hidden content: {e}")

#
# ------------------------ PDF STEGANOGRAPHY ------------------------

def pdf_to_binary(pdf_path):
    with open(pdf_path, 'rb') as file:
        binary_data = file.read()
    return binary_data

def embed_pdf_in_another(pdf_to_hide_path, base_pdf_path, output_pdf_path):
    try:
        binary_data = pdf_to_binary(pdf_to_hide_path)
        encoded_data = base64.b64encode(binary_data).decode('utf-8')

        with open(base_pdf_path, 'rb') as base_pdf:
            reader = PyPDF2.PdfReader(base_pdf)
            writer = PyPDF2.PdfWriter()

            for page_num in range(len(reader.pages)):
                writer.add_page(reader.pages[page_num])

            writer.add_metadata({
                '/CustomBinaryData': encoded_data
            })

            with open(output_pdf_path, 'wb') as output_pdf:
                writer.write(output_pdf)

        messagebox.showinfo("Success", f"Embedded PDF successfully in {output_pdf_path}")
    except Exception as e:
        messagebox.showerror("Error", f"Failed to hide PDF: {e}")

def extract_binary_from_pdf(pdf_path):
    try:
        with open(pdf_path, 'rb') as pdf_file:
            reader = PyPDF2.PdfReader(pdf_file)
            metadata = reader.metadata
            if '/CustomBinaryData' in metadata:
                encoded_data = metadata['/CustomBinaryData']
                binary_data = base64.b64decode(encoded_data)
                return binary_data
            else:
                raise ValueError("No embedded data found in the PDF metadata")
    except Exception as e:
        messagebox.showerror("Error", f"Failed to extract PDF: {e}")

def save_binary_as_pdf(binary_data, output_pdf_path):
    try:
        with open(output_pdf_path, 'wb') as output_pdf:
            output_pdf.write(binary_data)
        messagebox.showinfo("Success", f"Retrieved PDF successfully saved as {output_pdf_path}")
    except Exception as e:
        messagebox.showerror("Error", f"Failed to save the extracted PDF: {e}")

# ------------------------ DOCX STEGANOGRAPHY ------------------------

def docx_to_binary(docx_path):
    with open(docx_path, 'rb') as file:
        binary_data = file.read()
    return binary_data

def embed_docx_in_another(docx_to_hide_path, base_docx_path, output_docx_path):
    try:
        binary_data = docx_to_binary(docx_to_hide_path)
        encoded_data = base64.b64encode(binary_data).decode('utf-8')

        # Open the base DOCX document
        base_doc = Document(base_docx_path)

        # Create a hidden paragraph to embed the encoded data
        hidden_paragraph = base_doc.add_paragraph()
        hidden_paragraph.add_run(encoded_data).font.hidden = True  # Mark it as hidden

        # Save the modified DOCX document
        base_doc.save(output_docx_path)
        messagebox.showinfo("Success", f"Embedded DOCX successfully in {output_docx_path}")
    except Exception as e:
        messagebox.showerror("Error", f"Failed to hide DOCX: {e}")

def extract_binary_from_docx(docx_path):
    try:
        # Open the DOCX document
        doc = Document(docx_path)

        # Look for hidden paragraphs containing encoded data
        for para in doc.paragraphs:
            if para.runs and para.runs[0].font.hidden:
                encoded_data = para.runs[0].text
                binary_data = base64.b64decode(encoded_data)
                return binary_data

        raise ValueError("No embedded data found in the DOCX document")
    except Exception as e:
        messagebox.showerror("Error", f"Failed to extract DOCX: {e}")

# ------------------------ UI ------------------------

class SteganographyApp(tk.Tk):
    def init(self):
        super().init()
        self.title("Steganography Tool (Image, PDF, DOCX)")
        self.geometry("700x500")
        self.resizable(True, True)

        self.cover_image_path = tk.StringVar()
        self.secret_image_path = tk.StringVar()
        self.encoded_image_path = tk.StringVar()
        self.pdf_to_hide_path = tk.StringVar()
        self.base_pdf_path = tk.StringVar()
        self.encoded_pdf_path = tk.StringVar()
        self.docx_to_hide_path = tk.StringVar()
        self.base_docx_path = tk.StringVar()
        self.encoded_docx_path = tk.StringVar()

        self.create_sidebar()
        self.create_main_frames()

    def create_sidebar(self):
        sidebar = ttk.Frame(self, width=150, relief="ridge")
        sidebar.pack(side="left", fill="y")

        ttk.Label(sidebar, text="Steganography Options", font=("Arial", 12, "bold")).pack(pady=10)
        ttk.Button(sidebar, text="File Type", command=lambda: self.show_frame("file_type_page")).pack(fill="x", pady=5)
        ttk.Button(sidebar, text="Image Steganography", command=lambda: self.show_frame("image_page")).pack(fill="x", pady=5)
        ttk.Button(sidebar, text="PDF Steganography", command=lambda: self.show_frame("pdf_page")).pack(fill="x", pady=5)
        ttk.Button(sidebar, text="DOCX Steganography", command=lambda: self.show_frame("docx_page")).pack(fill="x", pady=5)

    def create_main_frames(self):
        self.frames = {}

        # File Type Page
        file_type_page = ttk.Frame(self)
        ttk.Label(file_type_page, text="Select File Type", font=("Arial", 14)).pack(pady=20)
        ttk.Button(file_type_page, text="Image", command=lambda: self.show_frame("image_page")).pack(pady=10)
        ttk.Button(file_type_page, text="PDF", command=lambda: self.show_frame("pdf_page")).pack(pady=10)
        ttk.Button(file_type_page, text="DOCX", command=lambda: self.show_frame("docx_page")).pack(pady=10)
        self.frames["file_type_page"] = file_type_page

        # Image Page
        image_page = ttk.Frame(self)
        ttk.Label(image_page, text="Image Steganography", font=("Arial", 14)).pack(pady=20)
        ttk.Button(image_page, text="Hide Image", command=lambda: self.show_frame("hide_image_page")).pack(pady=10)
        ttk.Button(image_page, text="Retrieve Image", command=lambda: self.show_frame("retrieve_image_page")).pack(pady=10)
        ttk.Button(image_page, text="Predict Hidden Content", command=self.predict_image_content).pack(pady=10)
        self.frames["image_page"] = image_page

        # Hide Image Page
        hide_image_page = ttk.Frame(self)
        ttk.Button(hide_image_page, text="Select Cover Image", command=self.select_cover_image).pack()
        ttk.Button(hide_image_page, text="Select Secret Image", command=self.select_secret_image).pack()
        ttk.Button(hide_image_page, text="Hide Image", command=self.hide_image_ui).pack(pady=10)
        self.frames["hide_image_page"] = hide_image_page

        # Retrieve Image Page
        retrieve_image_page = ttk.Frame(self)
        ttk.Button(retrieve_image_page, text="Select Encoded Image", command=self.select_encoded_image).pack()
        ttk.Button(retrieve_image_page, text="Retrieve Image", command=self.extract_image_ui).pack(pady=10)
        self.frames["retrieve_image_page"] = retrieve_image_page

        # PDF Page
        pdf_page = ttk.Frame(self)
        ttk.Label(pdf_page, text="PDF Steganography", font=("Arial", 14)).pack(pady=20)
        ttk.Button(pdf_page, text="Hide PDF", command=lambda: self.show_frame("hide_pdf_page")).pack(pady=10)
        ttk.Button(pdf_page, text="Retrieve PDF", command=lambda: self.show_frame("retrieve_pdf_page")).pack(pady=10)
        ttk.Button(pdf_page, text="Predict Hidden Content", command=self.predict_pdf_content).pack(pady=10)
        self.frames["pdf_page"] = pdf_page

        # Hide PDF Page
        hide_pdf_page = ttk.Frame(self)
        ttk.Button(hide_pdf_page, text="Select PDF to Hide", command=self.select_pdf_to_hide).pack()
        ttk.Button(hide_pdf_page, text="Select Base PDF", command=self.select_base_pdf).pack()
        ttk.Button(hide_pdf_page, text="Hide PDF", command=self.hide_pdf_ui).pack(pady=10)
        self.frames["hide_pdf_page"] = hide_pdf_page

        # Retrieve PDF Page
        retrieve_pdf_page = ttk.Frame(self)
        ttk.Button(retrieve_pdf_page, text="Select Encoded PDF", command=self.select_encoded_pdf).pack()
        ttk.Button(retrieve_pdf_page, text="Retrieve PDF", command=self.extract_pdf_ui).pack(pady=10)
        self.frames["retrieve_pdf_page"] = retrieve_pdf_page

        # DOCX Page
        docx_page = ttk.Frame(self)
        ttk.Label(docx_page, text="DOCX Steganography", font=("Arial", 14)).pack(pady=20)
        ttk.Button(docx_page, text="Hide DOCX", command=lambda: self.show_frame("hide_docx_page")).pack(pady=10)
        ttk.Button(docx_page, text="Retrieve DOCX", command=lambda: self.show_frame("retrieve_docx_page")).pack(pady=10)
        ttk.Button(docx_page, text="Predict Hidden Content", command=self.predict_docx_content).pack(pady=10)
        self.frames["docx_page"] = docx_page

        # Hide DOCX Page
        hide_docx_page = ttk.Frame(self)
        ttk.Button(hide_docx_page, text="Select DOCX to Hide", command=self.select_docx_to_hide).pack()
        ttk.Button(hide_docx_page, text="Select Base DOCX", command=self.select_base_docx).pack()
        ttk.Button(hide_docx_page, text="Hide DOCX", command=self.hide_docx_ui).pack(pady=10)
        self.frames["hide_docx_page"] = hide_docx_page

        # Retrieve DOCX Page
        retrieve_docx_page = ttk.Frame(self)
        ttk.Button(retrieve_docx_page, text="Select Encoded DOCX", command=self.select_encoded_docx).pack()
        ttk.Button(retrieve_docx_page, text="Retrieve DOCX", command=self.extract_docx_ui).pack(pady=10)
        self.frames["retrieve_docx_page"] = retrieve_docx_page

        self.show_frame("file_type_page")

    def show_frame(self, page_name):
        frame = self.frames[page_name]
        frame.pack(fill="both", expand=True)
        for f in self.frames.values():
            if f != frame:
                f.pack_forget()
    def predict_image_content(self):
        try:
            # Let the user choose an image file
            image_path = filedialog.askopenfilename(filetypes=[("Image Files", ".png;.jpg;.jpeg;.bmp")])
            if not image_path:
                return

            # Load the image and convert it to an array
            img = Image.open(image_path).convert("RGB")
            img_array = np.array(img)

            # Calculate the absolute differences between neighboring pixels
            diff_x = np.abs(img_array[:, 1:, :] - img_array[:, :-1, :])
            diff_y = np.abs(img_array[1:, :, :] - img_array[:-1, :, :])

            # Calculate the mean and standard deviation of these differences
            mean_diff = (np.mean(diff_x) + np.mean(diff_y)) / 2
            std_diff = (np.std(diff_x) + np.std(diff_y)) / 2

            # Set a dynamic threshold based on mean and standard deviation
            threshold = mean_diff + 2 * std_diff

            # Calculate the number of "anomalous" pixels based on the threshold
            anomalies = (diff_x > threshold).sum() + (diff_y > threshold).sum()

            # Determine if anomalies exceed an acceptable range (empirical factor for sensitivity)
            if anomalies > 0.05 * img_array.size:  # 5% of total pixels as a rough threshold
                messagebox.showinfo("Prediction", "Hidden content detected in the image!")
            else:
                messagebox.showinfo("Prediction", "No hidden content detected.")
        except Exception as e:
            messagebox.showerror("Error", f"Failed to predict hidden content in the image: {e}")

    def predict_pdf_content(self):
        try:
            pdf_path = filedialog.askopenfilename(filetypes=[("PDF Files", "*.pdf")])
            if not pdf_path:
                return

            with open(pdf_path, 'rb') as pdf_file:
                reader = PyPDF2.PdfReader(pdf_file)
                metadata = reader.metadata

                # Check for custom metadata entries as an indicator of hidden content
                if any(key.startswith('/Custom') for key in metadata.keys()):
                    messagebox.showinfo("Prediction", "Hidden content detected in the PDF!")
                else:
                    messagebox.showinfo("Prediction", "No hidden content detected.")
        except Exception as e:
            messagebox.showerror("Error", f"Failed to predict hidden content in the PDF: {e}")

    def predict_docx_content(self):
        try:
            docx_path = filedialog.askopenfilename(filetypes=[("DOCX Files", "*.docx")])
            if not docx_path:
                return

            doc = Document(docx_path)
            
            # Search for hidden text
            hidden_content_detected = any(
                para.runs and para.runs[0].font.hidden for para in doc.paragraphs
            )

            if hidden_content_detected:
                messagebox.showinfo("Prediction", "Hidden content detected in the DOCX!")
            else:
                messagebox.showinfo("Prediction", "No hidden content detected.")
        except Exception as e:
            messagebox.showerror("Error", f"Failed to predict hidden content in the DOCX: {e}")

    def select_cover_image(self):
        self.cover_image_path.set(filedialog.askopenfilename(filetypes=[("Image Files", ".png;.jpg;.jpeg;.bmp")])) 

    def select_secret_image(self):
        self.secret_image_path.set(filedialog.askopenfilename(filetypes=[("Image Files", ".png;.jpg;.jpeg;.bmp")])) 

    def hide_image_ui(self):
        output_image_path = filedialog.asksaveasfilename(defaultextension=".png", filetypes=[("PNG Files", "*.png")])
        if output_image_path:
            hide_image(self.cover_image_path.get(), self.secret_image_path.get(), output_image_path)

    def select_encoded_image(self):
        self.encoded_image_path.set(filedialog.askopenfilename(filetypes=[("Image Files", "*.png")])) 

    def extract_image_ui(self):
        output_secret_image_path = filedialog.asksaveasfilename(defaultextension=".png", filetypes=[("PNG Files", "*.png")])
        if output_secret_image_path:
            extract_image(self.encoded_image_path.get(), output_secret_image_path)

    def select_pdf_to_hide(self):
        self.pdf_to_hide_path.set(filedialog.askopenfilename(filetypes=[("PDF Files", "*.pdf")])) 

    def select_base_pdf(self):
        self.base_pdf_path.set(filedialog.askopenfilename(filetypes=[("PDF Files", "*.pdf")])) 

    def hide_pdf_ui(self):
        output_pdf_path = filedialog.asksaveasfilename(defaultextension=".pdf", filetypes=[("PDF Files", "*.pdf")])
        if output_pdf_path:
            embed_pdf_in_another(self.pdf_to_hide_path.get(), self.base_pdf_path.get(), output_pdf_path)

    def select_encoded_pdf(self):
        self.encoded_pdf_path.set(filedialog.askopenfilename(filetypes=[("PDF Files", "*.pdf")])) 

    def extract_pdf_ui(self):
        output_pdf_path = filedialog.asksaveasfilename(defaultextension=".pdf", filetypes=[("PDF Files", "*.pdf")])
        if output_pdf_path:
            binary_data = extract_binary_from_pdf(self.encoded_pdf_path.get())
            save_binary_as_pdf(binary_data, output_pdf_path)

    def select_docx_to_hide(self):
        self.docx_to_hide_path.set(filedialog.askopenfilename(filetypes=[("DOCX Files", "*.docx")])) 

    def select_base_docx(self):
        self.base_docx_path.set(filedialog.askopenfilename(filetypes=[("DOCX Files", "*.docx")])) 

    def hide_docx_ui(self):
        output_docx_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("DOCX Files", "*.docx")])
        if output_docx_path:
            embed_docx_in_another(self.docx_to_hide_path.get(), self.base_docx_path.get(), output_docx_path)

    def select_encoded_docx(self):
        self.encoded_docx_path.set(filedialog.askopenfilename(filetypes=[("DOCX Files", "*.docx")])) 

    def extract_docx_ui(self):
        output_docx_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("DOCX Files", "*.docx")])
        if output_docx_path:
            binary_data = extract_binary_from_docx(self.encoded_docx_path.get())
            with open(output_docx_path, 'wb') as output_file:
                output_file.write(binary_data)
            messagebox.showinfo("Success", f"Retrieved DOCX successfully saved as {output_docx_path}")
if name == "main":
    app = SteganographyApp()
    app.mainloop()





  
